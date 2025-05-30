import os
import streamlit as st
from custom_loaders.pdf_loader import SimplePDFLoader
from custom_loaders.text_loader import SimpleTextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from dotenv import load_dotenv
from deepseek_client import DeepSeekClient
from ollama_client import OllamaClient
import pandas as pd
import shutil
import time
import chromadb  # 添加ChromaDB导入
import config   # 导入配置文件

# 创建自定义DocxLoader处理Word文档
import docx
from langchain_core.documents import Document

# 添加在现有imports下面
from chromadb.api.client import EmbeddingFunction
import requests
import json
import tempfile
import re
import uuid
from langchain.prompts import PromptTemplate
from openai import OpenAI

# 添加Silicon Flow嵌入函数实现 - 适配ChromaDB 1.0.x API
class SiliconFlowEmbeddingFunction(EmbeddingFunction):
    def __init__(self, api_key, model_name="BAAI/bge-m3", api_url=None):
        self.api_key = api_key
        self.model_name = model_name
        self.api_url = api_url or "https://api.siliconflow.cn/v1"
        self.batch_size = 64  # Silicon Flow API批处理限制
    
    def __call__(self, input):
        """生成文本嵌入向量 - 支持批处理"""
        try:
            if not input:
                return []
            
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            # 实现批处理逻辑
            all_embeddings = []
            total_items = len(input)
            
            print(f"处理嵌入请求: 共{total_items}个文本，批处理大小: {self.batch_size}")
            
            # 将输入分成批次处理
            for i in range(0, total_items, self.batch_size):
                batch = input[i:i+self.batch_size]
                batch_num = i // self.batch_size + 1
                print(f"处理第{batch_num}批，包含{len(batch)}个文本")
                
                # 构建请求体
                data = {
                    "input": batch,
                    "model": self.model_name
                }
                
                # 发送请求
                response = requests.post(
                    f"{self.api_url}/embeddings", 
                    headers=headers, 
                    json=data,
                    timeout=60
                )
                
                # 如果请求成功
                if response.status_code == 200:
                    result = response.json()
                    # 提取嵌入向量
                    batch_embeddings = [item["embedding"] for item in result["data"]]
                    all_embeddings.extend(batch_embeddings)
                    print(f"第{batch_num}批嵌入生成成功，获取{len(batch_embeddings)}个向量")
                else:
                    print(f"第{batch_num}批嵌入请求失败，状态码：{response.status_code}")
                    print(f"错误详情：{response.text}")
                    # 为批次中的每个文本创建空向量
                    empty_embeddings = [[0.0] * 1024] * len(batch)
                    all_embeddings.extend(empty_embeddings)
            
            print(f"嵌入处理完成，共生成{len(all_embeddings)}个向量")
            return all_embeddings
                
        except Exception as e:
            print(f"生成嵌入向量时出错: {str(e)}")
            # 返回空向量列表
            return [[0.0] * 1024] * len(input)

class DocxLoader:
    def __init__(self, file_path):
        self.file_path = file_path
        
    def load(self):
        doc = docx.Document(self.file_path)
        text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        metadata = {"source": self.file_path}
        return [Document(page_content=text, metadata=metadata)]

# 创建自定义MarkdownLoader处理Markdown文件
class MarkdownLoader:
    def __init__(self, file_path):
        self.file_path = file_path
        
    def load(self):
        with open(self.file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        # 简单清理Markdown标记
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            # 移除标题符号
            if line.startswith('#'):
                line = line.lstrip('#').strip()
            # 移除链接
            while '[' in line and '](' in line and ')' in line:
                start = line.find('[')
                middle = line.find('](', start)
                end = line.find(')', middle)
                if start < middle < end:
                    link_text = line[start+1:middle]
                    line = line[:start] + link_text + line[end+1:]
                else:
                    break
            # 移除粗体和斜体
            line = line.replace('**', '').replace('__', '').replace('*', '').replace('_', '')
            cleaned_lines.append(line)
        
        cleaned_text = '\n'.join(cleaned_lines)
        metadata = {"source": self.file_path}
        return [Document(page_content=cleaned_text, metadata=metadata)]

# 加载环境变量
load_dotenv()

# 设置页面标题
st.title("个人知识库助手")

# 创建或获取会话状态中的LLM客户端
def get_llm_client():
    """根据当前会话状态获取或创建LLM客户端"""
    service_name = st.session_state.get("selected_service", "ollama")
    model_name = st.session_state.get(f"{service_name}_model", None)
    
    # 如果存在缓存的客户端且配置相同，直接返回
    cached_client = st.session_state.get("llm_client", None)
    cached_service = st.session_state.get("llm_service_name", None)
    cached_model = st.session_state.get("llm_model_name", None)
    
    if cached_client and cached_service == service_name and cached_model == model_name:
        return cached_client
    
    # 否则创建新客户端
    if service_name == "deepseek":
        # DeepSeek API客户端
        service_config = config.LLM_SERVICES["deepseek"]
        client = DeepSeekClient(timeout=config.APP_SETTINGS["request_timeout"])
        print(f"已创建DeepSeek客户端，API基础URL: {service_config['api_base']}")
    else:
        # Ollama API客户端
        service_config = config.LLM_SERVICES["ollama"]
        client = OllamaClient(
            api_base=service_config["api_base"],
            timeout=config.APP_SETTINGS["request_timeout"]
        )
        if model_name:
            client.set_model(model_name)
        else:
            client.set_model(service_config["default_model"])
        print(f"已创建Ollama客户端，使用模型: {client.model}")
    
    # 缓存客户端信息
    st.session_state["llm_client"] = client
    st.session_state["llm_service_name"] = service_name
    st.session_state["llm_model_name"] = model_name
    
    return client

# 侧边栏配置：模型选择
with st.sidebar:
    st.header("模型配置")
    
    # 语言模型服务选择
    available_services = []
    if config.LLM_SERVICES["deepseek"]["enabled"]:
        available_services.append("deepseek")
    if config.LLM_SERVICES["ollama"]["enabled"]:
        available_services.append("ollama")
    
    # 如果没有service_name选项，设置默认值
    if "selected_service" not in st.session_state:
        # 优先使用Ollama(如果可用)
        if "ollama" in available_services:
            st.session_state.selected_service = "ollama"
        else:
            st.session_state.selected_service = available_services[0] if available_services else "ollama"
    
    # 服务选择器
    service_options = {
        "deepseek": f"{config.LLM_SERVICES['deepseek']['name']} ({config.LLM_SERVICES['deepseek']['description']})",
        "ollama": f"{config.LLM_SERVICES['ollama']['name']} ({config.LLM_SERVICES['ollama']['description']})"
    }
    
    selected_service = st.selectbox(
        "选择语言模型服务",
        options=available_services,
        format_func=lambda x: service_options.get(x, x),
        key="selected_service"
    )
    
    # 根据所选服务显示相应的模型选项
    if selected_service == "deepseek":
        st.info("DeepSeek模型将使用配置文件中的API密钥")
    elif selected_service == "ollama":
        # 只有在确实选择了ollama服务时才初始化客户端
        try:
            # 获取Ollama可用模型
            ollama_client = OllamaClient(api_base=config.LLM_SERVICES["ollama"]["api_base"])
            available_models = ollama_client.get_available_models()
        except Exception as e:
            print(f"初始化Ollama客户端失败: {e}")
            available_models = []
        
        # 如果API获取不到模型，则使用配置文件中的默认列表
        if not available_models:
            available_models = config.LLM_SERVICES["ollama"]["models"]
        
        # 确保默认值在可用列表中
        default_model = config.LLM_SERVICES["ollama"]["default_model"]
        if default_model not in available_models and available_models:
            # 如果默认模型不可用，使用第一个可用模型
            default_model = available_models[0]
            print(f"配置的默认模型 {config.LLM_SERVICES['ollama']['default_model']} 不可用，使用 {default_model} 代替")
        
        # 设置默认模型
        if "ollama_model" not in st.session_state:
            st.session_state.ollama_model = default_model
        elif st.session_state.ollama_model not in available_models and available_models:
            # 如果会话状态中的模型不可用，更新为可用模型
            st.session_state.ollama_model = default_model
        
        # 模型选择器
        if available_models:
            selected_model = st.selectbox(
                "选择Ollama模型",
                options=available_models,
                key="ollama_model"
            )
            
            # 显示模型信息
            st.info(f"使用本地Ollama服务中的 {selected_model} 模型")
        else:
            st.error("未检测到可用的Ollama模型，请先安装并启动Ollama服务")
            # 创建一个空的选择框，避免错误
            st.selectbox("选择Ollama模型", ["请先安装Ollama模型"], disabled=True)
    
    # 显示嵌入模型信息
    st.header("嵌入模型")
    st.info(f"使用 {config.EMBEDDING_SERVICE['name']} 嵌入模型: {config.EMBEDDING_SERVICE['model']}")
    
    # 文档上传
    st.header("上传文档")
    uploaded_files = st.file_uploader(
        "选择PDF、TXT、Markdown或Word文档",
        accept_multiple_files=True,
        type=["pdf", "txt", "md", "docx", "doc"]
    )
    
    process_button = st.button("处理文档")

# 文档处理函数
def process_documents(uploaded_files):
    """处理上传的文档并创建知识库"""
    if not uploaded_files:
        st.sidebar.error("请上传至少一个文档")
        return

    with st.sidebar:
        with st.spinner("处理文档中..."):
    # 创建临时目录
            temp_dir = tempfile.mkdtemp()
            try:
                documents = []
                processed_files = []  # 跟踪成功处理的文件
                
                # 保存上传的文件并加载文档
                for uploaded_file in uploaded_files:
                    # 确保文件名是安全的
                    safe_filename = re.sub(r'[^\w\.-]', '_', uploaded_file.name)
                    temp_file_path = os.path.join(temp_dir, safe_filename)
                    
                    # 标准化路径格式（始终使用正斜杠）
                    norm_file_path = temp_file_path.replace("\\", "/")
                    
                    # 保存文件到临时目录
                    with open(temp_file_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())
                    
                    try:
                        # 根据文件类型处理文档
                        file_extension = os.path.splitext(safe_filename)[1].lower()
                        
                        if file_extension == '.pdf':
                            try:
                                print(f"\n开始处理PDF文件: {safe_filename}")
                                # 首先尝试使用SimplePDFLoader
                                loader = SimplePDFLoader(temp_file_path)
                                docs = loader.load()
                                
                                if not docs or not any(doc.page_content.strip() for doc in docs):
                                    print("SimplePDFLoader无法提取内容，尝试使用PyPDF直接提取")
                                    # 如果SimplePDFLoader未能提取到内容，尝试直接使用PyPDF
                                    import pypdf
                                    
                                    pdf_text = []
                                    with open(temp_file_path, "rb") as f:
                                        pdf = pypdf.PdfReader(f)
                                        for i, page in enumerate(pdf.pages):
                                            text = page.extract_text()
                                            if text and text.strip():
                                                pdf_text.append(text)
                                                print(f"PyPDF直接提取 - 页面 {i+1}: 提取了 {len(text)} 字符")
                                            else:
                                                print(f"PyPDF直接提取 - 页面 {i+1}: 无文本内容")
                                    
                                    if pdf_text:
                                        # 创建新的Document对象
                                        docs = [Document(
                                            page_content="\n\n".join(pdf_text),
                                            metadata={"source": norm_file_path, "pdf_fallback": True}
                                        )]
                                        print(f"使用PyPDF直接提取成功，获取了 {len(docs[0].page_content)} 字符")
                                    else:
                                        print("所有提取方法都失败，无法从PDF中获取文本")
                                        docs = []
                                else:
                                    print(f"SimplePDFLoader成功提取了 {len(docs)} 页内容")
                                    # 输出前10个字符作为预览
                                    for i, doc in enumerate(docs):
                                        if doc.page_content:
                                            preview = doc.page_content[:50].replace('\n', ' ') + "..."
                                            print(f"页面 {i+1} 预览: {preview}")
                            except Exception as e:
                                print(f"PDF处理异常: {str(e)}")
                                raise
                        elif file_extension == '.txt':
                            with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                text = f.read()
                            docs = [Document(page_content=text, metadata={"source": norm_file_path, "file_type": "text"})]
                        elif file_extension in ['.md', '.markdown']:
                            with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                text = f.read()
                            # 简单处理Markdown，移除代码块
                            text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
                            docs = [Document(page_content=text, metadata={"source": norm_file_path, "file_type": "markdown"})]
                        elif file_extension in ['.docx', '.doc']:
                            try:
                                from docx import Document as DocxDocument
                                doc = DocxDocument(temp_file_path)
                                text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                                docs = [Document(page_content=text, metadata={"source": norm_file_path, "file_type": "docx"})]
                            except ImportError:
                                st.error("未安装python-docx库，无法处理Word文档")
                                continue
                        else:
                            st.error(f"不支持的文件类型: {file_extension}")
                            continue
                        
                        # 统一元数据标准化处理
                        for doc in docs:
                            if not hasattr(doc, 'metadata') or not doc.metadata:
                                doc.metadata = {}
                                
                            # 确保关键元数据字段存在且格式一致
                            doc.metadata['source'] = norm_file_path  # 统一使用正斜杠
                            doc.metadata['filename'] = safe_filename
                            doc.metadata['file_type'] = file_extension[1:]  # 移除前导点
                            doc.metadata['doc_id'] = str(uuid.uuid4())  # 为每个文档分配唯一ID
                            doc.metadata['timestamp'] = int(time.time())
                            
                            # 对不同类型的文档添加特定元数据
                            if 'page' not in doc.metadata and file_extension == '.pdf':
                                doc.metadata['page'] = 1  # 默认页码
                        
                        documents.extend(docs)
                        processed_files.append(safe_filename)
                        st.success(f"成功处理: {safe_filename}")
                    
                    except Exception as e:
                        st.error(f"处理文件 {safe_filename} 时出错: {str(e)}")
                
                if not documents:
                    st.error("没有成功处理任何文档")
                    return
                
                # 分割文档
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=500,
                    chunk_overlap=50,
                    separators=["\n\n", "\n", "。", "！", "？", ".", " ", ""]
                )
                
                # 分块处理
                texts = []
                metadatas = []
                
                for doc in documents:
                    chunks = text_splitter.split_text(doc.page_content)
                    for i, chunk in enumerate(chunks):
                        if chunk.strip():  # 忽略空块
                            # 为每个文本块复制完整元数据
                            metadata = doc.metadata.copy()
                            # 添加块特定元数据
                            metadata['chunk'] = i
                            metadata['chunk_id'] = f"{metadata.get('doc_id', 'unknown')}_{i}"
                            metadata['total_chunks'] = len(chunks)
                            
                            texts.append(chunk)
                            metadatas.append(metadata)
                
                if not texts:
                    st.error("文档分割后没有有效内容")
                    return
                
                # 使用配置文件中的嵌入模型设置
                sf_embeddings = SiliconFlowEmbeddingFunction(
                    api_key=config.EMBEDDING_SERVICE["api_key"],
                    model_name=config.EMBEDDING_SERVICE["model"]
                )
                
                # 使用ChromaDB 1.0.6 API创建客户端和集合
                try:
                    # 创建ChromaDB客户端
                    client = chromadb.PersistentClient(path="./chroma_db")
                    
                    # 集合管理策略: 确保每个收集版本一致、状态正常
                    collection_name = "documents"
                    collection_metadata = {
                        "hnsw:space": "cosine",
                        "version": "1.0.1",  # 集合版本用于跟踪兼容性
                        "created_at": int(time.time()),
                        "last_updated": int(time.time())
                    }
                    
                    # 获取或创建集合(不再删除已有集合)
                    try:
                        # 尝试获取现有集合
                        collection = client.get_collection(
                            name=collection_name,
                            embedding_function=sf_embeddings
                        )
                        
                        # 更新集合元数据，保持版本一致性
                        current_metadata = collection.metadata
                        if current_metadata:
                            # 只更新last_updated字段，保留其他元数据
                            current_metadata["last_updated"] = int(time.time())
                            # 可以考虑在这里进行版本检查和兼容性处理
                        
                        print(f"获取到现有文档集合，包含 {collection.count()} 个文档")
                    except Exception as e:
                        print(f"集合不存在，将创建新集合: {str(e)}")
                        # 创建新集合
                        collection = client.create_collection(
                            name=collection_name,
                            embedding_function=sf_embeddings,
                            metadata=collection_metadata
                        )
                        print("创建了新的文档集合")
                    
                    # 准备文档ID - 使用真正唯一的ID而不仅仅是时间戳
                    # 格式: {时间戳}_{唯一标识符}_{索引}
                    timestamp = int(time.time())
                    batch_id = str(uuid.uuid4())[:8]  # 使用UUID的前8位作为批次标识
                    ids = [f"{timestamp}_{batch_id}_{i}" for i in range(len(texts))]
                    
                    # 将文档添加到集合中 - 使用批处理避免超时
                    batch_size = 16  # 批处理大小限制
                    total_batches = (len(texts) + batch_size - 1) // batch_size
                    
                    print(f"开始批量添加文档，共 {len(texts)} 个文本块，分 {total_batches} 批处理")
                    
                    for i in range(total_batches):
                        start_idx = i * batch_size
                        end_idx = min(start_idx + batch_size, len(texts))
                        
                        batch_texts = texts[start_idx:end_idx]
                        batch_metadatas = metadatas[start_idx:end_idx]
                        batch_ids = ids[start_idx:end_idx]
                        
                        print(f"添加第 {i+1}/{total_batches} 批文档，包含 {len(batch_texts)} 个文本块")
                        collection.add(
                            documents=batch_texts,
                            metadatas=batch_metadatas,
                            ids=batch_ids
                        )
                    
                    # 打印部分文档内容预览用于验证
                    print(f"\n====== 文档处理完成，共添加 {len(texts)} 个文档块 ======")
                    if texts:
                        # 随机打印3个文档块预览
                        import random
                        sample_indices = random.sample(range(len(texts)), min(3, len(texts)))
                        for idx in sample_indices:
                            preview = texts[idx][:150] + "..." if len(texts[idx]) > 150 else texts[idx]
                            source = metadatas[idx].get("source", "未知")
                            filename = metadatas[idx].get("filename", os.path.basename(source))
                            print(f"\n文档块 {idx+1}/{len(texts)} (来自 {filename}):")
                            print(f"内容预览: {preview}")
                            print(f"元数据: {metadatas[idx]}")
                    
                    st.success(f"已成功添加 {len(texts)} 个文档块到知识库")
                    
                    # 验证最新添加的文档
                    for id_to_check in [ids[0], ids[-1]]:
                        doc_check = collection.get(ids=[id_to_check], include=["metadatas", "documents"])
                        if doc_check and doc_check["metadatas"] and doc_check["documents"]:
                            print(f"验证文档ID [{id_to_check}]：")
                            print(f"- 元数据: {doc_check['metadatas'][0]}")
                            content_preview = doc_check['documents'][0][:50] + "..." if len(doc_check['documents'][0]) > 50 else doc_check['documents'][0]
                            print(f"- 内容预览: {content_preview}")
                        else:
                            print(f"警告: 无法验证文档ID [{id_to_check}]")
                
                except Exception as e:
                    st.error(f"创建知识库时出错: {str(e)}")
                    raise
            
            finally:
                # 清理临时目录
                try:
                    shutil.rmtree(temp_dir)
                except Exception as e:
                    print(f"清理临时目录时出错: {str(e)}")

# 查询知识库函数
def query_knowledge_base(question, document_filters=None):
    """查询知识库并生成回答"""
    try:
        # 获取ChromaDB集合，使用st.session_state缓存以避免重复获取
        if "chroma_client" not in st.session_state:
            print("初始化ChromaDB客户端并缓存到session_state")
            st.session_state.chroma_client = chromadb.PersistentClient(path="./chroma_db")
        
        if "chroma_collection" not in st.session_state:
            print("初始化ChromaDB集合并缓存到session_state")
            try:
                embedding_function = SiliconFlowEmbeddingFunction(
                    api_key=config.EMBEDDING_SERVICE["api_key"],
                    model_name=config.EMBEDDING_SERVICE["model"]
                )
                
                st.session_state.chroma_collection = st.session_state.chroma_client.get_collection(
                    name="documents",
                    embedding_function=embedding_function
                )
                print(f"成功获取集合，包含 {st.session_state.chroma_collection.count()} 个文档")
            except Exception as e:
                st.error(f"获取集合时出错: {str(e)}")
                st.error("知识库可能为空或损坏。请先上传文档。")
                return None
        
        # 使用缓存的集合
        collection = st.session_state.chroma_collection
        
        # 检测并处理复合问题
        sub_questions = split_compound_question(question)
        
        # 获取严格模式和复合问题检测设置
        strict_filter_mode = st.session_state.get("strict_filter_mode", True)
        disable_compound = st.session_state.get("disable_compound", False)

        print(f"查询使用严格过滤模式: {strict_filter_mode}, 禁用复合问题检测: {disable_compound}")

        # 如果禁用了复合问题检测，强制使用单一问题处理
        if disable_compound and len(sub_questions) > 1:
            print(f"用户选择禁用复合问题检测，将{len(sub_questions)}个子问题合并为1个处理")
            sub_questions = [question]  # 使用原始问题

        # 使用进度条显示处理状态
        progress_placeholder = st.empty()
        
        # 复合问题处理
        if len(sub_questions) > 1:
            print(f"检测到复合问题，拆分为 {len(sub_questions)} 个子问题")
            
            with progress_placeholder.container():
                st.write("正在处理复合问题...")
                progress_bar = st.progress(0)
            
            for i, sub_q in enumerate(sub_questions):
                print(f"子问题 {i+1}: {sub_q}")
                # 更新进度条
                progress_bar.progress((i / len(sub_questions)) * 0.5)  # 前50%用于问题检索
            
            # 为每个子问题查询相关文档
            all_documents = []
            all_metadatas = []
            all_ids = []
            
            for i, sub_q in enumerate(sub_questions):
                # 查询每个子问题 - 传递严格模式信息
                print(f"处理子问题 {i+1} '{sub_q}' (严格模式: {strict_filter_mode}):")
                # 确保process_single_question会用到严格模式标志
                st.session_state.strict_filter_mode = strict_filter_mode
                
                # 在严格模式下且有文档过滤时，明确提示
                if strict_filter_mode and document_filters:
                    print(f"为子问题'{sub_q}'应用严格文档过滤")
                    
                sub_results = process_single_question(collection, sub_q, document_filters)
                
                if not sub_results or not sub_results.get("documents") or not sub_results["documents"][0]:
                    print(f"子问题 '{sub_q}' 未找到匹配文档")
                    continue
                    
                # 从每个子问题获取最相关的文档（最多2个）
                if "documents" in sub_results and sub_results["documents"] and sub_results["documents"][0]:
                    max_docs = min(2, len(sub_results["documents"][0]))
                    for j in range(max_docs):
                        doc_content = sub_results["documents"][0][j]
                        doc_metadata = sub_results["metadatas"][0][j] if "metadatas" in sub_results and sub_results["metadatas"] and sub_results["metadatas"][0] else {}
                        doc_id = sub_results["ids"][0][j] if "ids" in sub_results and sub_results["ids"] and sub_results["ids"][0] else f"doc_{len(all_documents)}"
                        
                        # 打印文档来源和子问题的对应关系，用于调试
                        source = os.path.basename(doc_metadata.get("source", "未知")) if "source" in doc_metadata else "未知来源"
                        print(f"子问题 '{sub_q}' 找到匹配文档: {source} (ID: {doc_id})")
                        
                        all_documents.append(doc_content)
                        all_metadatas.append(doc_metadata)
                        all_ids.append(doc_id)
                
                # 更新进度条
                progress_bar.progress(0.5 + (i+1) / len(sub_questions) * 0.2)  # 50-70%用于组装结果
            
            # 构建结果字典
            if all_documents:
                results = {
                    "documents": [all_documents],
                    "metadatas": [all_metadatas],
                    "ids": [all_ids]
                }
                print(f"复合问题查询完成，合并得到 {len(all_documents)} 个相关文档")
            else:
                progress_placeholder.empty()
                print("所有子问题查询均未返回结果")
                return "很抱歉，在选定的文档中未找到与您问题相关的信息。"
        else:
            with progress_placeholder.container():
                st.write("正在查询知识库...")
                progress_bar = st.progress(0.3)  # 30%
            
            # 处理单一问题 
            results = process_single_question(collection, question, document_filters)
            progress_bar.progress(0.7)  # 70%
        
        # 处理查询结果
        if not results or not results["documents"] or not results["documents"][0]:
            progress_placeholder.empty()
            st.warning("没有找到相关文档，无法生成回答")
            return "很抱歉，在选定的文档中未找到与您问题相关的信息。"
        
        # 构建上下文
        contexts = []
        total_length = 0
        MAX_CONTEXT_LENGTH = config.APP_SETTINGS["max_context_length"]
        
        # 首先计算上下文总长度
        all_contexts = []
        for i, (doc, metadata) in enumerate(zip(results["documents"][0], results["metadatas"][0])):
            if metadata and "source" in metadata:
                source = os.path.basename(metadata["source"])
                context = f"[文档{i+1}] {source}:\n{doc}\n"
            else:
                context = f"[文档{i+1}]:\n{doc}\n"
            all_contexts.append((context, len(context)))
        
        # 按长度排序，优先保留较短的上下文
        all_contexts.sort(key=lambda x: x[1])
        
        # 添加上下文，确保总长度不超过限制
        for context, length in all_contexts:
            if total_length + length <= MAX_CONTEXT_LENGTH:
                contexts.append(context)
                total_length += length
            else:
                # 如果至少有一个上下文，就停止添加
                if contexts:
                    print(f"已达到上下文长度限制({total_length}/{MAX_CONTEXT_LENGTH})，剩余文档将被忽略")
                    break
                else:
                    # 如果一个上下文都没有，添加第一个并截断
                    truncated_context = context[:MAX_CONTEXT_LENGTH - 100] + "...[内容过长，已截断]"
                    contexts.append(truncated_context)
                    print(f"文档过长，已截断至{len(truncated_context)}字符")
                    break
        
        context_text = "\n".join(contexts)
        print(f"最终上下文长度: {len(context_text)}字符")
        
        progress_bar.progress(0.8)  # 80%
        
        # 构建提示
        prompt = PromptTemplate(
            template="""基于以下参考文档回答问题。如果文档中没有足够的信息，请说明无法回答。

参考文档:
{context}

问题: {question}

回答:""",
            input_variables=["context", "question"]
        )
        
        final_prompt = prompt.format(context=context_text, question=question)
        
        # 使用LLM生成回答
        with progress_placeholder.container():
            st.write("正在生成回答...")
            progress_bar.progress(0.9)  # 90%
        
        try:
            # 获取当前选择的LLM客户端
            llm_client = get_llm_client()
            answer = llm_client.ask(
                final_prompt, 
                temperature=config.APP_SETTINGS["temperature"], 
                timeout=config.APP_SETTINGS["request_timeout"]
            )
            
            # 检查是否返回错误消息
            if answer and answer.startswith("错误："):
                st.error(answer)
                answer = "很抱歉，生成回答时遇到了问题。请稍后再试。"
            
            progress_bar.progress(1.0)  # 100%
            progress_placeholder.empty()  # 完成后移除进度条
            return answer
        except Exception as e:
            progress_placeholder.empty()
            st.error(f"生成回答时出错: {str(e)}")
            return "很抱歉，生成回答时遇到了问题。请检查语言模型设置并稍后再试。"
    
    except Exception as e:
        st.error(f"查询知识库时出错: {str(e)}")
        return "很抱歉，查询知识库时发生了错误。请稍后再试。"

def split_compound_question(question):
    """拆分复合问题为多个子问题"""
    # 标记问题中可能的分隔符
    delimiters = ['？', '?', '；', ';']
    question_text = question.strip()
    
    # 检查是否存在多个分隔符
    multiple_delimiters = False
    delimiter_count = 0
    for d in delimiters[:2]:  # 只检查问号数量
        delimiter_count += question_text.count(d)
    
    multiple_delimiters = delimiter_count > 1
    
    if multiple_delimiters:
        print(f"检测到复合问题，包含{delimiter_count}个问号")
        
        # 统一问号和分号格式
        unified = question_text.replace('?', '？').replace(';', '；')
        
        # 先按问号分割
        if '？' in unified:
            parts = unified.split('？')
            # 过滤空字符串并重新添加问号
            sub_questions = [part.strip() + '？' for part in parts if part.strip()]
        # 如果没有问号但有分号，则按分号分割
        elif '；' in unified:
            parts = unified.split('；')
            # 过滤空字符串
            sub_questions = [part.strip() for part in parts if part.strip()]
        else:
            sub_questions = [question_text]  # 没有任何分隔符
        
        # 调试输出
        print(f"问题已拆分为{len(sub_questions)}个子问题:")
        for i, q in enumerate(sub_questions):
            print(f"  子问题{i+1}: {q}")
        
        return sub_questions
    
    # 单一问题
    print("问题被识别为单个问题，不进行拆分")
    return [question_text]

def process_single_question(collection, question, document_filters=None):
    """处理单个问题的查询"""
    # 查询相关文档
    query_params = {
        "query_texts": [question],
        "n_results": 5,
        "include": ["documents", "metadatas"]
    }
    
    # 初始化过滤集合为None，表示不过滤
    filter_ids_set = None
    # 获取严格模式设置 (优先使用会话状态)
    strict_filter_mode = st.session_state.get("strict_filter_mode", True)
    
    # 记录详细查询信息，用于调试
    print(f"处理问题: '{question}' (严格模式: {strict_filter_mode}, 文档过滤器: {'有' if document_filters else '无'})")
    
    # 强制执行严格过滤模式，当用户选择了单个文档时
    if document_filters and strict_filter_mode:
        is_single_doc = isinstance(document_filters, str) or (
            isinstance(document_filters, list) and len(document_filters) == 1
        )
        if is_single_doc:
            print("启用严格过滤模式：仅搜索用户选择的文档")
    
    # 如果指定了文档过滤条件，添加过滤器
    if document_filters:
        if not isinstance(document_filters, list):
            document_filters = [document_filters]  # 转换单个路径为列表
            
        print(f"应用文档过滤: {len(document_filters)}个文档，严格模式: {strict_filter_mode}")
        print(f"过滤文档列表: {[os.path.basename(doc) for doc in document_filters]}")
        
        # 获取所有文档元数据，在应用层面进行过滤
        try:
            # 获取所有文档 - 不要在include中指定ids，它会自动返回
            all_docs = collection.get(include=["metadatas"])
            
            if all_docs and "ids" in all_docs and all_docs["ids"]:
                print(f"获取到 {len(all_docs['ids'])} 个文档，准备进行过滤")
                
                # 根据文件名进行过滤，找到匹配的文档ID
                matching_ids = []
                
                # 处理每个文档过滤器
                for document_filter in document_filters:
                    # 标准化路径格式
                    document_filter = document_filter.replace("\\", "/")
                    filename = os.path.basename(document_filter)
                    
                    print(f"处理过滤器: {filename}")
                    
                    # 遍历所有文档，查找匹配的
                    for i, metadata in enumerate(all_docs["metadatas"]):
                        if metadata and "source" in metadata:
                            doc_source = metadata["source"]
                            doc_filename = os.path.basename(doc_source)
                            
                            # 使用多种方式检查匹配
                            if (filename in doc_source or
                                filename in doc_filename or
                                os.path.splitext(filename)[0] in doc_source):
                                matching_ids.append(all_docs["ids"][i])
                
                # 去重
                matching_ids = list(set(matching_ids))
                
                if matching_ids:
                    print(f"找到 {len(matching_ids)} 个精确匹配文档")
                    filter_ids_set = set(matching_ids)
                elif not strict_filter_mode:
                    # 只有在非严格模式下才尝试宽松匹配
                    # 没找到精确匹配，尝试更宽松的匹配
                    print("没有找到精确匹配，尝试更宽松的匹配")
                    
                    for document_filter in document_filters:
                        document_filter = document_filter.replace("\\", "/")
                        filename = os.path.basename(document_filter)
                        name_parts = os.path.splitext(filename)[0].split('_')
                        
                        # 再次遍历尝试部分匹配
                        for i, metadata in enumerate(all_docs["metadatas"]):
                            if metadata and "source" in metadata:
                                doc_source = metadata["source"]
                                for part in name_parts:
                                    if part and len(part) > 1 and part in doc_source:
                                        matching_ids.append(all_docs["ids"][i])
                                        break
                    
                    # 去重
                    matching_ids = list(set(matching_ids))
                    
                    if matching_ids:
                        print(f"使用部分匹配找到 {len(matching_ids)} 个文档")
                        filter_ids_set = set(matching_ids)
                    else:
                        print("无法找到任何匹配文档，将搜索所有文档")
                        filter_ids_set = None
                else:
                    print("在严格过滤模式下没有找到匹配文档，返回空结果")
                    # 在严格模式下若没有匹配，返回空结果集
                    return {
                        "documents": [[]],
                        "metadatas": [[]],
                        "ids": [[]]
                    }
            else:
                print("没有获取到文档元数据，将搜索所有文档")
                filter_ids_set = None
                
        except Exception as e:
            print(f"设置过滤条件时出错: {e}")
            print("将搜索所有文档")
            filter_ids_set = None
    
    # 执行查询
    print(f"执行查询，参数: {query_params}")
    results = collection.query(**query_params)
    
    # 调试: 打印原始查询结果
    if results and "documents" in results and results["documents"] and results["documents"][0]:
        print(f"原始查询获取到 {len(results['documents'][0])} 个相关文档")
        # 打印前3个文档的前100个字符(或全部)
        for i, doc in enumerate(results["documents"][0][:3]):
            preview = doc[:100] + "..." if len(doc) > 100 else doc
            print(f"文档{i+1}前100字符: {preview}")
            if "metadatas" in results and results["metadatas"] and results["metadatas"][0]:
                metadata = results["metadatas"][0][i]
                source = os.path.basename(metadata.get("source", "未知")) if metadata and "source" in metadata else "未知"
                print(f"文档{i+1}元数据: {metadata}")
                print(f"文档{i+1}来源: {source}")
    else:
        print("原始查询没有匹配结果")
    
    # 如果需要，在应用层面过滤结果
    if filter_ids_set and results and "ids" in results and results["ids"] and results["ids"][0]:
        print(f"在应用层面过滤查询结果，使用{len(filter_ids_set)}个文档ID")
        filtered_documents = []
        filtered_metadatas = []
        filtered_distances = []
        filtered_ids = []
        
        for i, doc_id in enumerate(results["ids"][0]):
            if doc_id in filter_ids_set:
                filtered_documents.append(results["documents"][0][i])
                filtered_metadatas.append(results["metadatas"][0][i])
                if "distances" in results and results["distances"] and results["distances"][0]:
                    filtered_distances.append(results["distances"][0][i])
                filtered_ids.append(doc_id)
        
        # 如果有过滤结果
        if filtered_documents:
            print(f"过滤后剩余 {len(filtered_documents)} 个文档")
            # 打印前2个过滤后文档的预览
            for i, doc in enumerate(filtered_documents[:2]):
                preview = doc[:100] + "..." if len(doc) > 100 else doc
                print(f"过滤后文档{i+1}前100字符: {preview}")
                if i < len(filtered_metadatas):
                    metadata = filtered_metadatas[i]
                    source = os.path.basename(metadata.get("source", "未知")) if metadata and "source" in metadata else "未知"
                    print(f"过滤后文档{i+1}元数据: {metadata}")
                    print(f"过滤后文档{i+1}来源: {source}")
            
            # 替换原始结果
            results["documents"] = [filtered_documents]
            results["metadatas"] = [filtered_metadatas]
            if "distances" in results and results["distances"]:
                results["distances"] = [filtered_distances]
            results["ids"] = [filtered_ids]
        else:
            print("过滤后没有匹配的文档")
            if strict_filter_mode:
                # 在严格模式下，如果过滤后没有文档，返回空结果
                print("严格模式下，过滤后没有匹配文档，返回空结果")
                return {
                    "documents": [[]],
                    "metadatas": [[]],
                    "ids": [[]]
                }
    
    return results

# 获取知识库中的所有文档
@st.cache_data(ttl=300)  # 缓存5分钟，避免频繁刷新
def get_document_list():
    """获取知识库中的所有文档，使用缓存减少重复调用"""
    if not os.path.exists("./chroma_db"):
        return []
    
    try:
        # 使用session_state中缓存的客户端
        if "chroma_client" not in st.session_state:
            print("[文档列表] 初始化ChromaDB客户端")
            st.session_state.chroma_client = chromadb.PersistentClient(path="./chroma_db")
        
        docs = []
        sources = set()
        
        # 尝试获取集合
        try:
            # 使用已缓存的集合或创建新的
            if "chroma_collection" not in st.session_state:
                print("[文档列表] 初始化ChromaDB集合")
                ef = SiliconFlowEmbeddingFunction(
                    api_key=config.EMBEDDING_SERVICE["api_key"],
                    model_name=config.EMBEDDING_SERVICE["model"]
                )
                
                st.session_state.chroma_collection = st.session_state.chroma_client.get_collection(
                    name="documents", 
                    embedding_function=ef
                )
            
            collection = st.session_state.chroma_collection
            print(f"[文档列表] 获取文档集合，包含 {collection.count()} 个文档")
            
            # 查询所有文档的元数据
            results = collection.get(include=["metadatas"])
            
            if results and "metadatas" in results and results["metadatas"]:
                # 提取所有不重复的source路径
                for metadata in results["metadatas"]:
                    if metadata and "source" in metadata and metadata["source"]:
                        sources.add(metadata["source"])
                    elif metadata and "filename" in metadata and metadata["filename"]:
                        # 如果有文件名但没有完整路径，尝试构建路径
                        temp_path = os.path.join("temp_docs", metadata["filename"])
                        if os.path.exists(temp_path):
                            sources.add(temp_path)
            else:
                print("[文档列表] 警告: 没有找到文档元数据")
        except Exception as e:
            print(f"[文档列表] 获取集合失败: {e}")
        
        # 如果从元数据中找不到文档源，尝试从文件系统获取
        if not sources and os.path.exists("temp_docs"):
            for filename in os.listdir("temp_docs"):
                file_path = os.path.join("temp_docs", filename)
                if os.path.isfile(file_path):
                    sources.add(file_path)
        
        # 构建文档列表
        for source in sources:
            filename = os.path.basename(source)
            docs.append({"文件名": filename, "路径": source})
        
        return docs
    except Exception as e:
        print(f"[文档列表] 获取文档列表出错: {e}")
        return []

# 删除特定文档
def delete_document(document_path):
    try:
        import time  # 确保time模块在函数内可用
        import chromadb
        
        print(f"尝试删除文档: {document_path}")
        # 标准化路径格式，使用正斜杠
        document_path_normalized = document_path.replace("\\", "/")
        filename = os.path.basename(document_path)
        print(f"标准化路径: {document_path_normalized}, 文件名: {filename}")
        
        # 1. 先尝试从ChromaDB中删除
        success_db = False
        try:
            # 初始化ChromaDB客户端
            client = chromadb.PersistentClient(path="./chroma_db")
            
            # 设置嵌入函数
            ef = SiliconFlowEmbeddingFunction(
                api_key=config.EMBEDDING_SERVICE["api_key"],
                model_name=config.EMBEDDING_SERVICE["model"]
            )
            
            # 获取集合
            try:
                collection = client.get_collection("documents", embedding_function=ef)
                print(f"获取到集合，包含 {collection.count()} 个文档")
            except Exception as e:
                print(f"获取集合失败: {e}")
                return False
            
            # 1.1 方法1：通过文件名删除
            try:
                print(f"尝试方法1 - 通过filename字段删除: '{filename}'")
                # 首先查询匹配的文档
                results = collection.get(
                    where={"filename": filename}
                )
                
                if results and results["ids"] and len(results["ids"]) > 0:
                    # 找到匹配的文档，执行删除
                    print(f"找到 {len(results['ids'])} 个匹配文档，IDs: {results['ids'][:5]}...")
                    collection.delete(
                        ids=results["ids"]
                    )
                    print(f"删除了 {len(results['ids'])} 个文档")
                    success_db = True
                else:
                    print("方法1：未找到匹配文档")
            except Exception as e:
                print(f"方法1删除失败: {e}")
            
            # 1.2 方法2：通过source路径删除
            if not success_db:
                try:
                    print(f"尝试方法2 - 通过source字段删除: '{document_path_normalized}'")
                    # 查询匹配的文档
                    results = collection.get(
                        where={"source": document_path_normalized}
                    )
                    
                    if results and results["ids"] and len(results["ids"]) > 0:
                        print(f"找到 {len(results['ids'])} 个匹配文档，IDs: {results['ids'][:5]}...")
                        collection.delete(
                            ids=results["ids"]
                        )
                        print(f"删除了 {len(results['ids'])} 个文档")
                        success_db = True
                    else:
                        print("方法2：未找到匹配文档")
                except Exception as e:
                    print(f"方法2删除失败: {e}")
            
            # 1.3 方法3：部分匹配删除
            if not success_db:
                try:
                    print("尝试方法3 - 部分匹配删除")
                    # 获取所有文档
                    all_docs = collection.get()
                    
                    ids_to_delete = []
                    if all_docs and all_docs["metadatas"]:
                        for i, metadata in enumerate(all_docs["metadatas"]):
                            if metadata:
                                src = metadata.get("source", "")
                                fname = metadata.get("filename", "")
                                
                                # 检查各种匹配情况
                                if (filename in src or 
                                    filename == fname or 
                                    document_path in src or
                                    document_path_normalized in src):
                                    ids_to_delete.append(all_docs["ids"][i])
                        
                        if ids_to_delete:
                            print(f"方法3：找到 {len(ids_to_delete)} 个匹配文档")
                            collection.delete(ids=ids_to_delete)
                            print(f"删除了 {len(ids_to_delete)} 个文档")
                            success_db = True
                        else:
                            print("方法3：未找到匹配文档")
                except Exception as e:
                    print(f"方法3删除失败: {e}")
        except Exception as e:
            print(f"从ChromaDB删除失败: {e}")
        
        # 2. 尝试删除实际文件
        success_file = False
        if os.path.exists(document_path):
            try:
                os.remove(document_path)
                print(f"文件删除成功: {document_path}")
                success_file = True
            except Exception as e:
                print(f"文件删除失败: {e}")
        
        # 3. 如果是临时目录中的文件，尝试重建知识库
        if (document_path.startswith("temp_docs/") or document_path.startswith("temp_docs\\")) and not success_db:
            print("尝试通过重建知识库方式删除文档...")
            
            # 列出所有需要保留的文件
            keep_files = []
            for file in os.listdir("temp_docs"):
                file_path = os.path.join("temp_docs", file)
                if file_path != document_path and os.path.isfile(file_path):
                    keep_files.append(file_path)
            
            print(f"需要保留的文件: {keep_files}")
            
            # 备份并重建ChromaDB
            backup_dir = f"./chroma_db_backup_{int(time.time())}"
            try:
                if os.path.exists("./chroma_db"):
                    os.rename("./chroma_db", backup_dir)
                    print(f"已备份旧数据库到: {backup_dir}")
                
                # 重新处理所有保留的文件
                if keep_files:
                    all_documents = []
                    
                    # 加载所有需要保留的文件
                    for file_path in keep_files:
                        file_ext = os.path.splitext(file_path)[1].lower()
                        try:
                            if file_ext == ".pdf":
                                loader = SimplePDFLoader(file_path)
                                docs = loader.load()
                                all_documents.extend(docs)
                            elif file_ext == ".txt":
                                loader = SimpleTextLoader(file_path)
                                docs = loader.load()
                                all_documents.extend(docs)
                            elif file_ext == ".md":
                                loader = MarkdownLoader(file_path)
                                docs = loader.load()
                                all_documents.extend(docs)
                            elif file_ext in [".docx", ".doc"]:
                                loader = DocxLoader(file_path)
                                docs = loader.load()
                                all_documents.extend(docs)
                        except Exception as e:
                            print(f"重新加载文件 {file_path} 失败: {e}")
                    
                    # 如果有加载成功的文档，重建知识库
                    if all_documents:
                        # 创建ChromaDB客户端
                        client = chromadb.PersistentClient(
                            path="./chroma_db",
                            settings=chromadb.Settings(
                                anonymized_telemetry=False,
                                allow_reset=True
                            )
                        )
                        
                        # 设置嵌入函数
                        ef = SiliconFlowEmbeddingFunction(
                            api_key=config.EMBEDDING_SERVICE["api_key"],
                            model_name=config.EMBEDDING_SERVICE["model"]
                        )
                        
                        # 创建新集合
                        collection = client.create_collection(
                            name="documents", 
                            embedding_function=ef,
                            metadata={"hnsw:space": "cosine"}
                        )
                        
                        # 分块文档
                        text_splitter = RecursiveCharacterTextSplitter(
                            chunk_size=1000,
                            chunk_overlap=200,
                            length_function=len
                        )
                        chunks = text_splitter.split_documents(all_documents)
                        
                        # 准备批量添加
                        doc_texts = []
                        doc_metadatas = []
                        doc_ids = []
                        
                        for i, doc in enumerate(chunks):
                            doc_texts.append(doc.page_content)
                            
                            # 简化元数据
                            simple_metadata = {
                                "source": doc.metadata.get("source", "").replace("\\", "/"),
                                "filename": os.path.basename(doc.metadata.get("source", f"doc_{i}")),
                                "index": i
                            }
                            
                            if "page" in doc.metadata:
                                simple_metadata["page"] = doc.metadata["page"]
                            
                            doc_metadatas.append(simple_metadata)
                            doc_ids.append(f"doc_{i}")
                        
                        # 批量添加文档
                        batch_size = 16
                        total_batches = (len(doc_texts) + batch_size - 1) // batch_size
                        
                        for i in range(total_batches):
                            start_idx = i * batch_size
                            end_idx = min(start_idx + batch_size, len(doc_texts))
                            
                            batch_texts = doc_texts[start_idx:end_idx]
                            batch_metadatas = doc_metadatas[start_idx:end_idx]
                            batch_ids = doc_ids[start_idx:end_idx]
                            
                            collection.add(
                                documents=batch_texts,
                                metadatas=batch_metadatas,
                                ids=batch_ids
                            )
                        
                        print(f"成功重建知识库，包含 {len(chunks)} 个文档块")
                        success_db = True
                
                # 清理备份
                try:
                    import threading
                    def delayed_delete():
                        time.sleep(5)
                        try:
                            if os.path.exists(backup_dir):
                                import shutil
                                shutil.rmtree(backup_dir)
                                print(f"成功删除备份: {backup_dir}")
                        except Exception as e:
                            print(f"删除备份失败: {e}")
                    
                    threading.Thread(target=delayed_delete, daemon=True).start()
                except Exception:
                    pass
                
                return True
            except Exception as e:
                print(f"重建知识库失败: {e}")
                
                # 尝试恢复备份
                try:
                    if os.path.exists(backup_dir) and not os.path.exists("./chroma_db"):
                        os.rename(backup_dir, "./chroma_db")
                        print("已恢复备份数据库")
                except Exception:
                    pass
        
        return success_db or success_file
    except Exception as e:
        st.error(f"删除文档时出错：{str(e)}")
        print(f"delete_document函数出错: {e}")
        return False

# 重建知识库
def rebuild_knowledge_base():
    try:
        import gc
        import os
        import time
        import psutil
        import shutil
        import sqlite3
        import chromadb
        
        # 先尝试修复SQLite数据库结构问题
        if os.path.exists("./chroma_db/chroma.sqlite3"):
            try:
                print("尝试修复数据库结构...")
                # 备份原始数据库
                backup_file = f"./chroma_db/chroma.sqlite3.bak_{int(time.time())}"
                shutil.copy("./chroma_db/chroma.sqlite3", backup_file)
                print(f"已备份原始数据库到: {backup_file}")
                
                # 尝试连接并修复数据库
                conn = None
                try:
                    conn = sqlite3.connect("./chroma_db/chroma.sqlite3")
                    cursor = conn.cursor()
                    
                    # 检查collections表结构
                    cursor.execute("PRAGMA table_info(collections)")
                    columns = cursor.fetchall()
                    column_names = [col[1] for col in columns]
                    
                    # 检查是否缺少topic列
                    if "topic" not in column_names:
                        print("检测到缺少collections.topic列，尝试添加...")
                        cursor.execute("ALTER TABLE collections ADD COLUMN topic TEXT")
                        conn.commit()
                        print("成功添加collections.topic列")
                    
                    # 检查数据库完整性
                    cursor.execute("PRAGMA integrity_check")
                    integrity_result = cursor.fetchone()
                    print(f"数据库完整性检查结果: {integrity_result[0]}")
                    
                    return True if integrity_result[0] == "ok" else False
                except sqlite3.Error as sql_err:
                    print(f"SQLite操作出错: {sql_err}")
                finally:
                    if conn:
                        conn.close()
            except Exception as db_err:
                print(f"尝试修复数据库结构失败: {db_err}")
        
        # 尝试关闭所有可能的数据库连接
        if os.path.exists("./chroma_db"):
            # 1. 先尝试通过API关闭连接
            try:
                ef = SiliconFlowEmbeddingFunction(
                    api_key=config.EMBEDDING_SERVICE["api_key"],
                    model_name=config.EMBEDDING_SERVICE["model"]
                )
                
                client = chromadb.PersistentClient(path="./chroma_db")
                try:
                    collection = client.get_collection("documents", embedding_function=ef)
                    # 尝试显式删除集合
                    client.delete_collection("documents")
                    print("成功删除集合")
                except Exception:
                    print("获取或删除集合失败，但会继续尝试")
                
                # 尝试强制删除对象
                del client
            except Exception as e:
                print(f"尝试关闭Chroma连接失败，但会继续尝试: {e}")
            
            # 2. 强制垃圾回收
            gc.collect()
            
            # 3. 等待一小段时间确保连接关闭
            time.sleep(2)
            
            # 4. 检查占用chroma_db的进程
            process = psutil.Process(os.getpid())
            for handler in process.open_files():
                if 'chroma_db' in handler.path:
                    print(f"警告: 发现文件仍然打开: {handler.path}")
            
            # 5. 尝试重命名目录而不是直接删除
            try:
                old_db = "./chroma_db"
                backup_db = "./chroma_db_old_" + str(int(time.time()))
                os.rename(old_db, backup_db)
                print(f"已将旧数据库重命名为: {backup_db}")
                
                # 如果重命名成功，尝试在后台异步删除
                def delayed_delete():
                    time.sleep(5)  # 等待5秒
                    try:
                        if os.path.exists(backup_db):
                            shutil.rmtree(backup_db)
                            print(f"成功删除旧数据库: {backup_db}")
                    except Exception:
                        pass
                
                # 在后台线程中执行删除
                import threading
                threading.Thread(target=delayed_delete, daemon=True).start()
                
                return True
            except Exception as e:
                # 如果重命名失败，尝试直接删除
                st.error(f"无法重命名数据库目录: {str(e)}")
                st.warning("将尝试直接删除数据库...")
                
                try:
                    shutil.rmtree("./chroma_db")
                    return True
                except Exception as e2:
                    st.error(f"无法删除数据库目录: {str(e2)}")
                    st.info("请尝试手动关闭所有Streamlit进程后再试，或手动删除chroma_db目录")
                    st.code("taskkill /f /im streamlit.exe && rm -r -Force chroma_db")
                    return False
        
        return True  # 如果数据库不存在，视为重建成功
    except Exception as e:
        st.error(f"重建知识库时出错：{str(e)}")
        return False

# 添加知识库健康检查函数
def check_knowledge_base_health():
    """检查知识库健康状态并修复可能的问题"""
    try:
        if not os.path.exists("./chroma_db"):
            return {"status": "error", "message": "知识库不存在"}

        client = chromadb.PersistentClient(path="./chroma_db")
        
        # 1. 检查集合是否存在
        try:
            collections = client.list_collections()
            if not collections or "documents" not in [c.name for c in collections]:
                return {"status": "error", "message": "知识库中没有documents集合"}
                
            # 获取文档集合
            ef = SiliconFlowEmbeddingFunction(
                api_key=config.EMBEDDING_SERVICE["api_key"],
                model_name=config.EMBEDDING_SERVICE["model"]
            )
            
            collection = client.get_collection("documents", embedding_function=ef)
            doc_count = collection.count()
            
            if doc_count == 0:
                return {"status": "warning", "message": "知识库存在但为空", "count": 0}
            
            # 2. 检查随机文档的完整性
            try:
                # 获取所有IDs
                all_ids = collection.get(include=["metadatas"])["ids"]
                if not all_ids:
                    return {"status": "warning", "message": "知识库中没有文档ID", "count": 0}
                    
                # 随机抽查5个文档或全部(取较小值)
                sample_size = min(5, len(all_ids))
                import random
                sample_ids = random.sample(all_ids, sample_size)
                
                # 检查这些文档是否有完整的元数据和内容
                samples = collection.get(ids=sample_ids, include=["metadatas", "documents"])
                
                # 计算元数据和内容的完整性
                metadata_integrity = sum(1 for meta in samples["metadatas"] if meta and "source" in meta) / sample_size
                content_integrity = sum(1 for doc in samples["documents"] if doc and len(doc) > 10) / sample_size
                
                integrity_score = (metadata_integrity + content_integrity) / 2
                
                report = {
                    "status": "healthy" if integrity_score > 0.8 else "warning",
                    "message": f"知识库健康度: {integrity_score*100:.1f}%",
                    "count": doc_count,
                    "metadata_integrity": f"{metadata_integrity*100:.1f}%",
                    "content_integrity": f"{content_integrity*100:.1f}%",
                    "collections": [c.name for c in collections]
                }
                
                return report
            except Exception as check_err:
                return {"status": "warning", "message": f"健康检查失败: {str(check_err)}", "count": doc_count}
                
        except Exception as coll_err:
            return {"status": "error", "message": f"获取集合失败: {str(coll_err)}"}
            
    except Exception as e:
        return {"status": "error", "message": f"知识库健康检查失败: {str(e)}"}

# 处理文档上传
if process_button and uploaded_files:
    with st.spinner("处理文档中..."):
        process_documents(uploaded_files)

# 创建选项卡界面
tab1, tab2, tab3 = st.tabs(["问答", "知识库管理", "帮助"])

# 问答界面
with tab1:
    st.header("向知识库提问")
    
    # 获取可用文档列表
    available_docs = get_document_list()
    doc_names = ["全部文档"] + [doc["文件名"] for doc in available_docs]
    
    # 将单选改为多选
    selected_docs = st.multiselect("选择要查询的文档 (可多选)", doc_names, default=["全部文档"])
    
    # 添加严格模式开关
    strict_mode = st.checkbox("严格模式 (仅从选定文档中查询，不使用相似匹配)", value=True)
    
    # 添加禁用复合问题检测选项 - 新增
    disable_compound = st.checkbox("禁用复合问题检测 (多个问题一起查询时选择)", value=False, 
                                help="启用后，多个问题将作为一个整体查询，而不是分别查询每个问题")
    
    # 使用表单控制提交行为，只在按回车时提交
    with st.form("query_form"):
        query = st.text_input("输入您的问题")
        submit_button = st.form_submit_button("提交", type="primary")

        # 只在表单提交时处理查询
        if submit_button and query:
            if not os.path.exists("./chroma_db"):
                st.error("知识库为空，请先处理文档")
            elif not config.EMBEDDING_SERVICE["api_key"] or not config.EMBEDDING_SERVICE["model"]:
                if not config.EMBEDDING_SERVICE["api_key"]:
                    st.warning("请输入嵌入模型API密钥")
                if not config.EMBEDDING_SERVICE["model"]:
                    st.warning("请输入嵌入模型名称")
            else:
                with st.spinner("思考中..."):
                    # 确定文档筛选条件
                    doc_filters = None
                    if selected_docs and "全部文档" not in selected_docs:
                        # 找到对应的文档完整路径
                        doc_filters = []
                        for selected_doc in selected_docs:
                            for doc in available_docs:
                                if doc["文件名"] == selected_doc:
                                    doc_filters.append(doc["路径"])
                                    print(f"选择了文档: {selected_doc}, 过滤路径: {doc['路径']}")
                    
                    # 传递严格模式设置
                    st.session_state.strict_filter_mode = strict_mode
                    
                    # 传递复合问题检测设置
                    st.session_state.disable_compound = disable_compound
                    
                    # 获取回答
                    answer = query_knowledge_base(query, doc_filters)
                    
                    # 显示回答
                    st.markdown("### 回答")
                    st.write(answer)

# 知识库管理界面
with tab2:
    st.header("知识库管理")
    
    # 强制重建知识库选项
    st.subheader("知识库维护")
    st.warning("⚠️ 如果应用反复出错，可能是知识库版本不兼容导致的，请尝试重建知识库")
    
    if st.button("强制重建知识库", type="primary"):
        with st.spinner("正在重建知识库..."):
            try:
                if os.path.exists("./chroma_db"):
                    shutil.rmtree("./chroma_db")
                st.success("知识库已重建，请重新上传并处理文档")
                st.rerun()
            except Exception as e:
                st.error(f"重建知识库时出错：{str(e)}")
    
    # 知识库状态
    if os.path.exists("./chroma_db"):
        docs = get_document_list()
        if docs:
            st.success(f"知识库中包含 {len(docs)} 个文档")
            
            # 显示文档列表
            df = pd.DataFrame(docs)
            st.dataframe(df)
            
            # 删除文档功能
            st.subheader("删除文档")
            doc_to_delete = st.selectbox("选择要删除的文档", [doc["文件名"] for doc in docs])
            
            if st.button("删除所选文档"):
                # 找到文档路径
                doc_path = ""
                for doc in docs:
                    if doc["文件名"] == doc_to_delete:
                        doc_path = doc["路径"]
                        break
                
                if doc_path:
                    with st.spinner("正在删除文档..."):
                        if delete_document(doc_path):
                            st.success(f"已删除文档: {doc_to_delete}")
                            st.rerun()
        else:
            st.info("知识库中没有文档或知识库可能已损坏")
        
        # 正常重建知识库功能
        st.subheader("重建知识库")
        st.warning("⚠️ 此操作将删除所有向量数据，请谨慎操作!")
        if st.button("重建知识库"):
            with st.spinner("正在重建知识库..."):
                if rebuild_knowledge_base():
                    st.success("知识库已重建，请重新上传并处理文档")
                    st.rerun()
    else:
        st.info("知识库尚未创建，请先上传并处理文档")

    # 添加知识库健康检查功能
    st.subheader("知识库诊断")
    if st.button("运行健康检查"):
        with st.spinner("正在检查知识库健康状态..."):
            health_report = check_knowledge_base_health()
            
            if health_report["status"] == "healthy":
                st.success(f"✅ 知识库健康: {health_report['message']}")
            elif health_report["status"] == "warning":
                st.warning(f"⚠️ 知识库警告: {health_report['message']}")
            else:
                st.error(f"❌ 知识库错误: {health_report['message']}")
            
            # 显示详细报告
            if "count" in health_report:
                st.info(f"文档数量: {health_report['count']}")
            
            if "metadata_integrity" in health_report:
                cols = st.columns(2)
                cols[0].metric("元数据完整性", health_report["metadata_integrity"])
                cols[1].metric("内容完整性", health_report["content_integrity"])

# 帮助界面
with tab3:
    st.header("使用说明")
    st.markdown("""
    ### 个人知识库助手
    
    这是一个基于DeepSeek和Silicon Flow API的智能文档管理与问答系统。
    
    #### 基本功能
    
    1. **文档处理**：上传PDF、TXT、Markdown或Word文档，系统会自动提取文本并向量化存储
    2. **智能问答**：基于文档内容回答您的问题
    3. **文档管理**：查看、删除知识库中的文档
    
    #### 使用步骤
    
    1. 在侧边栏输入API密钥
    2. 上传文档并点击"处理文档"按钮
    3. 在"问答"选项卡输入问题，可以选择查询特定文档或全部文档
    4. 在"知识库管理"选项卡管理您的文档
    
    #### 注意事项
    
    - 首次使用需要上传并处理文档才能进行问答
    - API密钥不会被保存，每次启动应用都需要重新输入
    - 处理大型文档可能需要一定时间
    """)
