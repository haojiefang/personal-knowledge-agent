
import docx
from typing import List
from langchain.docstore.document import Document

class SimpleDocxLoader:
    """简化版Word文档加载器"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """加载Word文档内容"""
        try:
            doc = docx.Document(self.file_path)
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            metadata = {"source": self.file_path}
            return [Document(page_content=text, metadata=metadata)]
        except Exception as e:
            print(f"加载Word文档时出错: {e}")
            return []
